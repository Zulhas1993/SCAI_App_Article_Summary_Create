# Import necessary libraries
import os
from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential
import nltk
from newspaper import Article
from docx import Document
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.tag import pos_tag
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import random
from nltk.corpus import wordnet

# Set your Azure Text Analytics service credentials
# Replace "YOUR_TEXT_ANALYTICS_ENDPOINT" and "YOUR_TEXT_ANALYTICS_KEY" with your actual values
endpoint = "https://laboblogtextanalytics.cognitiveservices.azure.com/"
key = "09b7c88cfff44bac95c83a2256f31907"

# Function to authenticate Azure Text Analytics client
def authenticate_client():
    ta_credential = AzureKeyCredential(key)
    text_analytics_client = TextAnalyticsClient(endpoint=endpoint, credential=ta_credential)
    return text_analytics_client

# Function to detect language using Azure Text Analytics
def detect_language(text_analytics_client, texts):
    documents = [{"id": str(i+1), "text": text} for i, text in enumerate(texts)]
    result = text_analytics_client.detect_language(documents=documents)
    detected_languages = [res.primary_language.iso6391_name if res.primary_language else None for res in result]
    return detected_languages

# Function to analyze sentiment using Azure Text Analytics
def analyze_sentiment(text_analytics_client, texts):
    documents = [{"id": str(i+1), "language": "en", "text": text} for i, text in enumerate(texts)]
    result = text_analytics_client.analyze_sentiment(documents=documents)
    sentiments = [res.sentiment for res in result]
    return sentiments

# Function to extract key phrases using Azure Text Analytics
def extract_key_phrases(text_analytics_client, text):
    result = text_analytics_client.extract_key_phrases(text)
    return result.key_phrases

# Function to calculate cosine similarity between two texts

def calculate_cosine_similarity(text1, text2):
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([text1, text2])
    similarity = cosine_similarity(vectors[0], vectors[1])[0][0]
    return similarity

# Function to get synonyms of a word based on its part of speech
def get_synonyms(word, pos):
    synonyms = []
    for syn in wordnet.synsets(word, pos=pos):
        for lemma in syn.lemmas():
            synonyms.append(lemma.name())
    return synonyms

# Function to generate a short sentence with similar parts of speech
def generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords):
    parts_of_speech = {'N': 'noun', 'V': 'verb', 'JJ': 'adjective'}

    words_by_pos = {pos: [word for word, pos_tag in tagged_sentence if pos_tag.startswith(pos) and word.lower() not in user_request_keywords]
                    for pos in parts_of_speech}

    if not all(words_by_pos.values()):
        return None

    # Use placeholders for different parts of speech
    placeholders = {pos: random.choice(words) for pos, words in words_by_pos.items()}
    
    # Identify the context of the words and construct a more dynamic sentence
    sentence = construct_dynamic_sentence(tagged_sentence, placeholders, user_title_keywords, user_request_keywords)

    return sentence

# Function to construct a dynamic sentence based on context
def construct_dynamic_sentence(tagged_sentence, placeholders, user_title_keywords, user_request_keywords):
    # Placeholder for dynamic sentence construction based on context
    # You should replace this with your custom logic
    # Example: Use surrounding words to create a more coherent sentence
    return f"In {' '.join(user_title_keywords)}, {' '.join(user_request_keywords)}. {placeholders['N'][0]} {placeholders['V'][0]} {placeholders['JJ'][0]}."

# Function to generate a short summary of an article
# sent_tokenize:
# The sent_tokenize function is part of the NLTK (Natural Language Toolkit) library.
# It is used to tokenize a given text into a list of sentences.
# In this context, selected_scraped_data['details_text'] is assumed to contain the text of an article or document.
# word_tokenize:
# The word_tokenize function is another NLTK function used for tokenizing sentences into words.
# The word_tokenize function is applied to each sentence obtained from sent_tokenize.
# pos_tag:
# The pos_tag function is also part of the NLTK library and is used for part-of-speech tagging.
# Part-of-speech tagging involves assigning a part-of-speech category (such as noun, verb, adjective, etc.) to each word in a sentence.


def generate_short_summary(summary, user_title_keywords, user_request_keywords, target_word_count=160):
    sentences = sent_tokenize(summary)
    tagged_sentences = [pos_tag(word_tokenize(sentence)) for sentence in sentences]
    short_summary_content = []
    word_count = 0
    for tagged_sentence in tagged_sentences:
        new_sentence = generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords)
        if new_sentence:
            short_summary_content.append(new_sentence)
            word_count += len(word_tokenize(new_sentence))
            if word_count >= target_word_count:
                break

    return " ".join(short_summary_content)

# Function to generate a unique article based on user input and scraped data
def generate_unique_article(scraped_data, user_title_keywords, user_request_keywords):
    similarities = [
        0.8 * calculate_cosine_similarity(user_title, item['title']) +
        0.6 * calculate_cosine_similarity(user_request, item['details_text'])
        for item in scraped_data
    ]

    # Identify the index with the highest overall similarity
    max_similarity_index = similarities.index(max(similarities))

    # Use the data from the most relevant scraped item
    selected_scraped_data = scraped_data[max_similarity_index]

    sentences = sent_tokenize(selected_scraped_data['details_text'])
    tagged_sentences = [pos_tag(word_tokenize(sentence)) for sentence in sentences]

    new_article_content = []
    target_word_count = 140
    word_count = 0

    while word_count < target_word_count and tagged_sentences:
        tagged_sentence = tagged_sentences.pop(0)  # Use the first sentence
        new_sentence = generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords)

        if new_sentence:
            new_article_content.append(new_sentence)
            word_count += len(word_tokenize(new_sentence))

    return " ".join(new_article_content), selected_scraped_data['url']

# Function to authenticate Azure Text Analytics client, scrape articles, analyze and summarize them
def analyze_and_summarize_with_azure(article_urls, user_title, user_request, output_file="summarized_articles.docx", new_article_output_file="new_article.docx"):
    try:
        text_analytics_client = authenticate_client()
        doc = Document()
        scraped_data = []

        # Loop through the provided article URLs
        for article_url in article_urls:
            article = Article(article_url)
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
            article.headers = headers

            # Download and parse the article using the newspaper library
            article.download()
            article.parse()

            # Perform NLP tasks on the article
            nltk.download('punkt')
            article.nlp()

            # Extract title and details from the article
            title = article.title
            details = article.text

            # Detect language and analyze sentiment using Azure Text Analytics
            language = detect_language(text_analytics_client, [details])
            sentiment = analyze_sentiment(text_analytics_client, [details])

            # Store the scraped data in a list
            scraped_data.append({'title': title, 'details_text': details, 'url': article_url, 'language': language, 'sentiment': sentiment})

        # Extract user title and request keywords
        user_title_keywords = set(user_title.lower().split())
        user_request_keywords = set(user_request.lower().split())

        max_similarity = -1
        selected_scraped_data = None

        # Calculate similarity and find the most relevant scraped data
        for item in scraped_data:
            title_similarity = calculate_cosine_similarity(user_title, item['title'])
            details_similarity = calculate_cosine_similarity(user_request, item['details_text'])
            overall_similarity = 0.7 * title_similarity + 0.3 * details_similarity

            if overall_similarity > max_similarity:
                max_similarity = overall_similarity
                selected_scraped_data = item

        # Check if relevant data is found
        if selected_scraped_data:
            title_matched = any(keyword in selected_scraped_data['title'].lower() for keyword in user_title_keywords)
            request_matched = any(keyword in selected_scraped_data['details_text'].lower() for keyword in user_request_keywords)

            # Add information to the document
            doc.add_heading(selected_scraped_data['title'], level=1)

            doc.add_heading("User Interest Analysis", level=2)
            doc.add_paragraph(f"User Title Keywords: {', '.join(user_title_keywords)}")
            doc.add_paragraph(f"User Request Keywords: {', '.join(user_request_keywords)}")
            doc.add_paragraph(f"Title Matched: {'Yes' if title_matched else 'No'}")
            doc.add_paragraph(f"Request Matched: {'Yes' if request_matched else 'No'}")

            doc.add_heading("Details", level=2)
            doc.add_paragraph(selected_scraped_data['details_text'])

            # Generate short summary based on the matched content
            if title_matched or request_matched:
                short_summary_content = generate_short_summary(selected_scraped_data['details_text'], user_title_keywords, user_request_keywords)
            else:
                short_summary_content = generate_short_summary(selected_scraped_data['title'], user_title_keywords, user_request_keywords)

            doc.add_heading("Short Summary", level=2)
            doc.add_paragraph(short_summary_content)

            # Generate a unique article and save it to a new document
            try:
                new_article_content, relevant_url = generate_unique_article(scraped_data, user_title_keywords, user_request_keywords)
                if new_article_content:
                    new_article_doc = Document()
                    new_article_doc.add_heading(f"Unique Article based on '{user_title}'", level=1)
                    new_article_doc.add_paragraph(new_article_content)
                    new_article_doc.save(new_article_output_file)
                    print(f"The unique article has been saved to '{new_article_output_file}'. Relevant URL: {relevant_url}")
            except Exception as e:
                print(f"An error occurred while generating the unique article: {e}")

        # Save the analyzed and summarized articles to a document
        doc.save(output_file)
        print(f"The analyzed and summarized articles have been saved to '{output_file}'.")
    except Exception as ex:
        print(f"An error occurred: {ex}")

if __name__ == "__main__":
    # Example article URLs for testing
    article_urls = [
        "https://www.thedailystar.net/news/bangladesh/news/bangladesh-india-joint-water-measurement-padma-ganges-begins-3508056",
        "https://www.thedailystar.net/news/bangladesh/national-election-2024/news/cec-warns-crisis-if-anti-polls-activities-descend-violence-3508021",
        "https://www.thedailystar.net/news/bangladesh/transport/news/metro-rail-now-stops-shahbagh-karwan-bazar-stations-3507041",
        "https://www.nhk.or.jp/d-navi/note/article/20210430.html",
        "https://www.nhk.or.jp/d-navi/note/article/20230714.html",
    ]

    # Get user input for user title and request
    user_title = input("Enter user title: ")
    user_request = input("Enter user request: ")

    # Call the main function to analyze and summarize articles
    analyze_and_summarize_with_azure(article_urls, user_title, user_request)
