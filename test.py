import requests
from bs4 import BeautifulSoup
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from docx import Document

# Constants
SCRAPE_URL = "https://www.thedailystar.net/news/bangladesh/crime-justice/news/serial-muggers-terrorise-mohammadpur-3506776"
USER_INTEREST = {
    "title": "Mohammadpur Terrorism",
    "request": "I am a resident of Mohammadpur, deeply concerned about the rising incidents of muggings and robberies in our area. I want to know the latest updates and measures taken by authorities to address this issue."
}

# Function to scrape news feed
def scrape_news_feed(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Extracting news data
    title = soup.find('h1', class_='fw-700 e-mb-16 article-title').text.strip()

    # Modify the details extraction based on the actual HTML structure
    details = ""
    target_class = 'pb-20 clearfix'
    target_div = soup.find('div', class_=target_class)

    if target_div:
        paragraphs = target_div.find_all('p')
        details = '\n'.join(paragraph.text for paragraph in paragraphs)

    return [{'title': title, 'details': details}]

# Function to preprocess and tokenize text
def preprocess_text(text):
    stop_words = set(stopwords.words("english"))
    words = word_tokenize(text)
    words = [word.lower() for word in words if word.isalpha() and word.lower() not in stop_words]
    return words

# Function to create a summarized article
def create_summarized_article(news_data, user_interest):
    # Combine news data and user interest
    combined_data = f"{user_interest['request']} \n\n Latest News: \n"
    for news_item in news_data:
        combined_data += f"\nTitle: {news_item['title']}\nDetails: {news_item['details']}\n"

    # Tokenize and extract keywords using nltk
    all_words = preprocess_text(combined_data)
    freq_dist = FreqDist(all_words)
    keywords = freq_dist.most_common(5)  # Extract top 5 keywords (you can adjust this number)

    # Summarize the article based on keywords
    summarized_article = f"Keywords: {', '.join(word for word, _ in keywords)}\n\n{combined_data}"

    return summarized_article

# Function to save the article to a DOCX file
def save_to_docx(article, filename):
    doc = Document()
    doc.add_heading('Summarized Article', level=1)
    doc.add_paragraph(article)
    doc.save(filename)

def main():
    # Step 1: Scrape news feed
    scraping_data = scrape_news_feed(SCRAPE_URL)

    # Step 2: Create a summarized article
    summarized_article = create_summarized_article(scraping_data, USER_INTEREST)

    # Step 3: Save the summarized article to a DOCX file
    save_to_docx(summarized_article, 'summarized_article.docx')

if __name__ == "__main__":
    main()
