import nltk
from newspaper import Article
from docx import Document
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.tag import pos_tag
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import random

def calculate_cosine_similarity(text1, text2):
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([text1, text2])
    similarity = cosine_similarity(vectors[0], vectors[1])[0][0]
    return similarity

def generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords):
    # Identify nouns, verbs, and adjectives in the sentence
    nouns = [word for word, pos in tagged_sentence if pos.startswith('N') and word.lower() not in user_request_keywords]
    verbs = [word for word, pos in tagged_sentence if pos.startswith('V') and word.lower() not in user_request_keywords]
    adjectives = [word for word, pos in tagged_sentence if pos.startswith('JJ') and word.lower() not in user_request_keywords]

    # Check if any of the lists are non-empty
    if not nouns or not verbs or not adjectives:
        return None

    # Shuffle the order of word categories to create variation
    random.shuffle(nouns)
    random.shuffle(verbs)
    random.shuffle(adjectives)

    # Generate a new sentence with predefined sentence templates
    sentence_templates = [
        "In {}, {}. {} {} {} {}.",
        "While {}, {}. {} {} {} {}.",
        "Exploring {}, {}. {} {} {} {}.",
    ]

    new_sentence_template = random.choice(sentence_templates)
    new_sentence = new_sentence_template.format(
        random.choice(list(user_title_keywords)),
        random.choice(list(user_request_keywords)),
        random.choice(nouns),
        random.choice(verbs),
        random.choice(adjectives),
        random.choice(list(user_request_keywords)),
    )

    return new_sentence

def generate_short_summary(summary, user_title_keywords, user_request_keywords, target_word_count=160):
    # Tokenize the summary into sentences
    sentences = sent_tokenize(summary)

    # Perform part-of-speech tagging on each sentence
    tagged_sentences = [pos_tag(word_tokenize(sentence)) for sentence in sentences]

    # Create a short summary by generating sentences with similar parts of speech
    short_summary_content = []
    word_count = 0

    for tagged_sentence in tagged_sentences:
        new_sentence = generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords)
        if new_sentence:  # Check if a valid sentence is generated
            short_summary_content.append(new_sentence)
            word_count += len(word_tokenize(new_sentence))
            if word_count >= target_word_count:
                break

    return " ".join(short_summary_content)

def generate_unique_article(scraped_data, user_title_keywords, user_request_keywords):
    # Extract titles and details from the scraped data
    scraped_titles = [item['title'] for item in scraped_data]
    scraped_details = [item['details_text'] for item in scraped_data]

    # Calculate cosine similarity between user title/request and scraped titles/details
    title_similarities = [calculate_cosine_similarity(user_title, title) for title in scraped_titles]
    details_similarities = [calculate_cosine_similarity(user_request, details) for details in scraped_details]

    # Identify the index with the highest similarity for both title and details
    max_title_index = title_similarities.index(max(title_similarities))
    max_details_index = details_similarities.index(max(details_similarities))

    # Use the data from the most relevant scraped item
    selected_scraped_data = scraped_data[max(max_title_index, max_details_index)]

    # Tokenize the details into sentences
    sentences = sent_tokenize(selected_scraped_data['details_text'])

    # Perform part-of-speech tagging on each sentence
    tagged_sentences = [pos_tag(word_tokenize(sentence)) for sentence in sentences]

    # Create a new article content by generating sentences with similar parts of speech
    new_article_content = []
    target_word_count = 140
    word_count = 0

    while word_count < target_word_count:
        for tagged_sentence in tagged_sentences:
            new_sentence = generate_short_sentence_with_similar_pos(tagged_sentence, user_title_keywords, user_request_keywords)
            if new_sentence:  # Check if a valid sentence is generated
                new_article_content.append(new_sentence)
                word_count += len(word_tokenize(new_sentence))

    return " ".join(new_article_content)

def analyze_and_summarize(article_urls, user_title, user_request, output_file="summarized_articles.docx", new_article_output_file="new_article.docx"):
    # Create a Document object for the summary
    doc = Document()

    # Scraping data storage
    scraped_data = []

    for article_url in article_urls:
        # Download and parse the article
        article = Article(article_url)
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        article.headers = headers

        article.download()
        article.parse()

        # Perform natural language processing
        nltk.download('punkt')
        article.nlp()

        # Extract article title, top image, and details
        title = article.title
        details = article.text

        # Save scraped data for later comparison
        scraped_data.append({'title': title, 'details_text': details, 'url': article_url})

    # Extract user interest keywords
    user_title_keywords = set(user_title.lower().split())
    user_request_keywords = set(user_request.lower().split())

    # Initialize variables to track the most relevant URL
    max_similarity = -1
    selected_scraped_data = None

    # Iterate through scraped data to find the most relevant URL
    for item in scraped_data:
        title_similarity = calculate_cosine_similarity(user_title, item['title'])
        details_similarity = calculate_cosine_similarity(user_request, item['details_text'])
        overall_similarity = 0.7 * title_similarity + 0.3 * details_similarity

        if overall_similarity > max_similarity:
            max_similarity = overall_similarity
            selected_scraped_data = item

    if selected_scraped_data:
        # Check if any user interest keyword is present in both the article title and request
        title_matched = any(keyword in selected_scraped_data['title'].lower() for keyword in user_title_keywords)
        request_matched = any(keyword in selected_scraped_data['details_text'].lower() for keyword in user_request_keywords)

        # Add title to the document
        doc.add_heading(selected_scraped_data['title'], level=1)

        # Add user interest information to the document
        doc.add_heading("User Interest Analysis", level=2)
        doc.add_paragraph(f"User Title Keywords: {', '.join(user_title_keywords)}")
        doc.add_paragraph(f"User Request Keywords: {', '.join(user_request_keywords)}")
        doc.add_paragraph(f"Title Matched: {'Yes' if title_matched else 'No'}")
        doc.add_paragraph(f"Request Matched: {'Yes' if request_matched else 'No'}")

        # Add details to the document
        doc.add_heading("Details", level=2)
        doc.add_paragraph(selected_scraped_data['details_text'])

        # Generate a short summary from the dynamic details
        if title_matched or request_matched:  # Use the article title if it matches the user title or request
            short_summary_content = generate_short_summary(selected_scraped_data['details_text'], user_title_keywords, user_request_keywords)
        else:  # Use the user title for generating a short summary
            short_summary_content = generate_short_summary(selected_scraped_data['title'], user_title_keywords, user_request_keywords)

        # Add short summary to the document
        doc.add_heading("Short Summary", level=2)
        doc.add_paragraph(short_summary_content)

        # Generate a unique article from the dynamic details
        new_article_content = generate_unique_article([selected_scraped_data], user_title_keywords, user_request_keywords)

        if new_article_content:
            # Create a Document object for the new article
            new_article_doc = Document()

            # Add title to the new article document
            new_article_doc.add_heading(f"Unique Article based on '{user_title}'", level=1)  # Use user title as the header

            # Add generated content to the new article document
            new_article_doc.add_paragraph(new_article_content)

            # Save the new article document to a separate file
            new_article_doc.save(new_article_output_file)
            print(f"The unique article has been saved to '{new_article_output_file}'. Relevant URL: {selected_scraped_data['url']}")

    # Save the document to the specified output file
    doc.save(output_file)
    print(f"The analyzed and summarized articles have been saved to '{output_file}'.")

if __name__ == "__main__":
    # Example usage: Replace the URLs with the ones you want to analyze and summarize
    article_urls = [
        "https://www.thedailystar.net/news/bangladesh/news/bangladesh-india-joint-water-measurement-padma-ganges-begins-3508056",
        "https://www.thedailystar.net/news/bangladesh/national-election-2024/news/cec-warns-crisis-if-anti-polls-activities-descend-violence-3508021",
        "https://www.thedailystar.net/news/bangladesh/transport/news/metro-rail-now-stops-shahbagh-karwan-bazar-stations-3507041",
        "https://www.nhk.or.jp/d-navi/note/article/20210430.html",
        "https://www.nhk.or.jp/d-navi/note/article/20230714.html",
    ]

    user_title = input("Enter user title: ")
    user_request = input("Enter user request: ")

    analyze_and_summarize(article_urls, user_title, user_request)
