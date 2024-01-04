import os
import csv
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from azure.cognitiveservices.language.textanalytics import TextAnalyticsClient
from msrest.authentication import CognitiveServicesCredentials
import datetime
from azure.cognitiveservices.language.textanalytics.models import DetectLanguageInput

# Constants
SCRAPE_URL = "https://b.hatena.ne.jp/"
CSV_FILE_PATH = "scraping_data.csv"
DOCX_FILE_PATH = "scraping_data.docx"
DB_FILE_PATH = "theme_info.csv"
AZURE_KEY = "5e1835fa2e784d549bb1b2f6bd6ed69f"
AZURE_ENDPOINT = "https://labo-azure-openai-swedencentral.openai.azure.com/"

# h3,a,href,title, class = entrylist-contents-main,entrylist-contents-title,
# span,class = entrylist-contents-users, entrylist-contents-domain,entrylist-contents-body,
# class = entrylist-contents-meta, entrylist-contents-tags,
# div, class=entry-content, hatenablog-entry, <p>
# div,  class= section




# Function to scrape news feed


# def scrape_news_feed():
#     response = requests.get(SCRAPE_URL)
#     soup = BeautifulSoup(response.text, 'html.parser')
    
#     # Extracting news data
#     news_list = []
#     for item in soup.find_all('h3', class_='entrylist-contents-title'):
#         title = item.a.text
#         href = item.a['href']
#         news_list.append({'title': title, 'href': href})
    
#     return news_list


def scrape_news_feed():
    response = requests.get(SCRAPE_URL)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Extracting news data
    news_list = []
    for item in soup.find_all('div', class_='news-item'):  # Adjust the HTML structure based on the actual website
        title = item.find('h3').text
        details_url = item.find('a')['href']
        details = scrape_news_details(details_url)  # Scrape details for each news item
        news_list.append({'title': title, 'details_url': details_url, 'details': details})

    return news_list

# Function to scrape details for each news item
def scrape_news_details(details_url):
    # Implement logic to scrape details from the news details page
    # You can use requests and BeautifulSoup to extract information from the details page
    # Return the details as a string or a structured format
    details_response = requests.get(details_url)
    details_soup = BeautifulSoup(details_response.text, 'html.parser')
    details = details_soup.find('div', class_='news-details').text  # Adjust based on the actual HTML structure
    return details






# Function to update CSV and DOCX files with scraping data
def update_files_with_scraping_data(data):
    # Update CSV file
    df = pd.DataFrame(data)
    if not os.path.exists(CSV_FILE_PATH):
        df.to_csv(CSV_FILE_PATH, index=False, header=True)
    else:
        df.to_csv(CSV_FILE_PATH, mode='a', index=False, header=False)

    # Update DOCX file
    doc = Document()
    for entry in data:
        doc.add_paragraph(f"Title: {entry['title']}\nHref: {entry['href']}\n")
    doc.save(DOCX_FILE_PATH)

# Function to fetch data from the database
def fetch_data_from_db():
    # Fetch data from the ScaiDb table
    # Assume you have a method to fetch data from the database
    # and return a list of dictionaries
    db_data = [{'title': 'Sample Title', 'request': 'Sample Request', 'user_id': 1}]  # Replace with actual data
    return db_data

# Function to update CSV file with database data
def update_file_with_db_data(data):
    df = pd.DataFrame(data)
    if not os.path.exists(DB_FILE_PATH):
        df.to_csv(DB_FILE_PATH, index=False, header=True)
    else:
        df.to_csv(DB_FILE_PATH, mode='a', index=False, header=False)

# Function to analyze data using Azure Text Analytics
def analyze_data_with_azure(data):
    credentials = CognitiveServicesCredentials(AZURE_KEY)
    text_analytics = TextAnalyticsClient(endpoint=AZURE_ENDPOINT, credentials=credentials)

    # Extracting themes from the data
    themes = []
    for entry in data:
        text = entry['request']
        theme = detect_themes(text_analytics, text)
        themes.append({'id': entry['id'], 'themes': theme})

    return themes

# Function to create articles using AI service
def create_articles(data, themes):
    # Assume you have a method to create articles using data and themes
    # and return a list of articles
    articles = [{'title': 'Article Title', 'content': 'Article Content'}]  # Replace with actual data
    return articles

# Function to summarize articles
def summarize_articles(articles):
    # Assume you have a method to summarize articles
    # and return a list of summarized articles
    summarized_articles = [{'title': 'Summarized Article Title', 'summary': 'Article Summary'}]  # Replace with actual data
    return summarized_articles

# Function to detect themes using Azure Text Analytics
def detect_themes(client, text):
    documents = [DetectLanguageInput(id='1', text=text)]
    response = client.detect_language(documents=documents)

    if response.documents[0].detected_languages[0].iso6391_name == 'en':
        # Language is English, extract key phrases
        key_phrases = client.key_phrases(documents=documents)
        return key_phrases.documents[0].key_phrases
    else:
        return []
    
def save_summarized_articles_to_docx(summarized_articles):
    doc = Document()
    doc.add_heading('Summarized Articles', level=1)
    
    for article in summarized_articles:
        doc.add_heading(article['title'], level=2)
        doc.add_paragraph(article['summary'])
        doc.add_paragraph("\n")  # Add some space between articles

    doc.save("summarized_articles.docx")    

# Main function
def main():
    # Step 1: Scrape news feed
    scraping_data = scrape_news_feed()

    # Step 2: Update CSV and DOCX files with scraping data
    update_files_with_scraping_data(scraping_data)

    # Step 3: Fetch data from the database
    db_data = fetch_data_from_db()

    # Step 4: Update CSV file with database data
    update_file_with_db_data(db_data)

    # Step 5: Analyze data with Azure Text Analytics
    themes = analyze_data_with_azure(db_data)

    # Step 6: Create articles using AI service
    articles = create_articles(scraping_data, themes)

    # Step 7: Summarize articles
    summarized_articles = summarize_articles(articles)

    # Step 8: Save summarized articles to DOCX file
    save_summarized_articles_to_docx(summarized_articles)

    # Print or save the summarized articles as needed
    for article in summarized_articles:
        print(f"Title: {article['title']}\nSummary: {article['summary']}\n")

    # Print or save the summarized articles as needed
    for article in summarized_articles:
        print(f"Title: {article['title']}\nSummary: {article['summary']}\n")

if __name__ == "__main__":
    main()