import nltk
from newspaper import Article
from docx import Document

def analyze_and_summarize(url, user_title, user_request, output_file="summarized_article.docx"):
    # Download and parse the article
    article = Article(url)
    article.download()
    article.parse()

    # Perform natural language processing
    nltk.download('punkt')
    article.nlp()

    # Extract article title, top image, and summary
    title = article.title
    summary = article.summary

    # Extract user interest keywords
    user_title_keywords = set(user_title.lower().split())
    user_request_keywords = set(user_request.lower().split())

    # Check if any user interest keyword is present in both the article title and request
    title_matched = any(keyword in title.lower() for keyword in user_title_keywords)
    request_matched = any(keyword in summary.lower() for keyword in user_request_keywords)

    # Create a Document object
    doc = Document()

    # Add title to the document
    doc.add_heading(title, level=1)

    # Add user interest information to the document
    doc.add_heading("User Interest Analysis", level=2)
    doc.add_paragraph(f"User Title Keywords: {', '.join(user_title_keywords)}")
    doc.add_paragraph(f"User Request Keywords: {', '.join(user_request_keywords)}")
    doc.add_paragraph(f"Title Matched: {'Yes' if title_matched else 'No'}")
    doc.add_paragraph(f"Request Matched: {'Yes' if request_matched else 'No'}")

    # Add summary to the document
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)

    # Save the document to the specified output file
    doc.save(output_file)
    print(f"The analyzed and summarized article has been saved to '{output_file}'.")

if __name__ == "__main__":
    # Example usage: Replace the URL with the one you want to analyze and summarize
    article_url = "https://www.thedailystar.net/news/bangladesh/crime-justice/news/serial-muggers-terrorise-mohammadpur-3506776"

    # Set user title and request dynamically
    user_title = input("Enter user title: ")
    user_request = input("Enter user request: ")

    analyze_and_summarize(article_url, user_title, user_request)
